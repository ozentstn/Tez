from turtledemo.round_dance import stop
import content
import nltk.data
import pypyodbc
import re
import os
import nltk
import string
import numpy as np
import itertools
import math
import datetime
import snowballstemmer
import spacy
import tokenize
from nltk.tokenize import word_tokenize,sent_tokenize
from nltk.probability import FreqDist
import docx
from nltk.tag import pos_tag
from bs4 import BeautifulSoup
from nltk.stem.porter import PorterStemmer
from nltk.corpus import stopwords
from nltk.corpus import wordnet as wn
from nltk.tokenize import word_tokenize
from nltk.tag import pos_tag
from string import punctuation
import matplotlib.pyplot as plt #kmeans
import pandas as pd #kmeans
# Using the elbow method to find the optimal number of clusters
from sklearn.cluster import KMeans
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer

#Read Word
"""file_path="C:\\Users\\ozent\\OneDrive\\Masaüstü\\orhanPamukYeniHayat.docx"
def read_docx(file_path):
    document = docx.Document(file_path)
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text)
    return 'n'.join(text)
document=read_docx(file_path)

#stringi array/series/df çevir
document_split=document.split(" ")
vector=pd.Series(document_split)
document_vector=vector[1:len(vector)] #0 dan başladığı için 1 den başlasın
document_df=pd.DataFrame(document_vector,columns=["kelimeler"])
#küçük harf dönüşümü
copy_document_df=document_df.copy() #orjinali dursun
copy_document_dff=pd.DataFrame(copy_document_df,columns=["kelimeler"])
upperLowerCase=copy_document_dff["kelimeler"].apply(lambda x:" ".join(x.lower() for x in x.split()))
#az kelimelerin silinmesi
upperLowerCaseDf=pd.DataFrame(upperLowerCase,columns=["kelimeler"])
delete=pd.Series(''.join(upperLowerCaseDf['kelimeler']).split()).value_counts()<'5'
azKelime=upperLowerCaseDf['kelimeler'].apply(lambda x: " ".join(x for x in x.split() if x not in delete))
#sayıların silinmesi
#upperLowerCaseDf=pd.DataFrame(upperLowerCase,columns=["kelimeler"])
removeNumbers=[x for x in azKelime if not x.isnumeric()]
#noktalama işaretleri
noktalama=str.maketrans('','',string.punctuation)
removeNoktalama=[x.translate(noktalama) for x in removeNumbers]
#stopwords silinmesi
sw=stopwords.words("turkish")
stopWords=[x for x in removeNoktalama if x not in sw]

#lemmatization: zemberek/ snowballStemmer : kök bulma
texts=[]
p_stemmer=PorterStemmer
tr_stemmer=snowballstemmer.stemmer('turkish')
for i in stopWords:
   stemmed=[tr_stemmer.stemWord(i) for i in stopWords if not i in sw]
   stemmed_tokens=[i for i in stemmed if not i in sw]
   #texts.append(stemmed_tokens)
"""
file_path="C:\\Users\\ozent\\OneDrive\\Masaüstü\\orhanPamukYeniHayat.docx"
def read_docx(file_path):
    document = docx.Document(file_path)
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text)
    return 'n'.join(text)
document=read_docx(file_path)

def getFreq(content):
    translator=str.maketrans('','',string.punctuation)
    #kelime ayrıştırma
    kelimeler=word_tokenize(content)
    kelimeler=[word.translate(translator) for word in kelimeler]
    kelimeler=[word for word in kelimeler if not word.isnumeric()]
    kelimeler=[word.lower() for word in kelimeler]
    kelimeler=[word for word in kelimeler if word not in stopwords.words('turkish')]
    #fdist = FreqDist(kelimeler)
    return kelimeler
kelimeler_df=pd.DataFrame(getFreq(document),columns=["Sütun1"])
fdist = FreqDist(kelimeler_df)
#ayrık değerlerin çıkarılması
print(fdist)

#zarf,sıfat
#nltk.download('averaged_perceptron_tagger')
#postag=pos_tag(etkisiz_kelimeler)












